#c_png2docx.py
#to be run under PYthon 2.7
from docx import Document
from docx.shared import Cm, RGBColor, Pt
from docx.text.run import Font, Run
from docx.shared import Inches, Mm
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
# from docx.dml.color import ColorFormat

document = Document()

fontn = document.styles['Normal'].font
fontn.size = Pt(9)
fontn.name = 'Frutiger LT 45 Light'
fontn.color.rgb = RGBColor(31, 73, 125)

document.add_heading('PHAST Analysis - ' + folder, level=0)

ISs = []
for e in lEvent:
    pvis,hole,weather = e.Key.split("\\")
    if not (pvis in ISs):
        ISs.append(pvis)

for IS in ISs:
    document.add_heading(IS, level=1)
    for e in lEvent:
        pvis,hole,weather = e.Key.split("\\")
        if IS == pvis:                        
            fn = 'C\\'+folder+'\\C_'+slugify(e.Key)+'_2.png'
            document.add_picture(fn, width=Cm(15))            

            p = document.add_paragraph('Figure',style='FigureCaption')
            Figure(p)
            p.add_run(pvis + " Hole: " + hole + " Weather: " + weather)            
    document.add_page_break()
document.save('C_'+folder+'_Rev.2.docx')