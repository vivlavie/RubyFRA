from openpyxl import load_workbook
import math
import numpy as np
import matplotlib.pyplot as plt
from scipy.interpolate import interp1d
import scipy.ndimage
from scipy.ndimage.filters import gaussian_filter
import dill

import matplotlib.pylab as pltparam
from matplotlib.patches import Ellipse

pltparam.rcParams["figure.figsize"] = (8,3)
pltparam.rcParams['lines.linewidth'] = 2
# pltparam.rcParams['lines.color'] = 'r'
pltparam.rcParams['axes.grid'] = True 


import re
def slugify(value):
    """
    Normalizes string, converts to lowercase, removes non-alpha characters,
    and converts spaces to hyphens.
    """
    # import unicodedata
    # value = unicodedata.normalize('NFKD', value).encode('ascii', 'ignore')
    # value = unicode(re.sub('[^\w\s-]', '', value).strip().lower())
    # value = unicode(re.sub('[-\s]+', '-', value))
    value = re.sub('[^\w\s-]', '', value).strip().lower()
    value = re.sub('[-\s]+', '-', value)
    return value

# element_dump_filename = 'Bv06_offloading_dump'
# folder = 'Offloading'
# XDir = -1 #From Stern to Bow

# element_dump_filename = 'Bv06_dump'
# folder = 'ProcessArea'
# XDir = 1 #From Bow to Stern

# element_dump_filename = 'Bv06_utility_dump'
element_dump_filename = 'Bv06_utility_dump_2' #with the increased leak frequency
folder = 'Utility'
# XDir = 1 #From Bow to Stern

# element_dump_filename = 'Bv08_c5_dump'
# folder = 'Process'
# XDir = 1 #From Bow to Stern


with open(element_dump_filename,'rb') as element_dump:
    lEvent = dill.load(element_dump)

pool_diameter = [0.5, 2.5, 5, 10, 25, 30, 50, 80,100] #the lowest pool diamter and its radiation distance is added by SHK refering to ther ation at 2.5m
distance_12  = [1.26, 6.3, 9.4, 11, 13.2, 15.6, 25.2, 40.2, 50]
distance_5 = [2.4, 12, 19, 28, 30, 32, 43, 64, 77]
#Distance from the center of pool i.e. radius
f12 = interp1d(pool_diameter,distance_12)
f5 = interp1d(pool_diameter,distance_5)

for e in lEvent:    
# for e in [lEvent[5]]:    
    # if ("-L" in e.Key) and ("LM_EO" in e.Hole) and (e.Weather == "7.7D"):    
        Key = e.Key
        X = e.X
        Y = e.Y        
        
        if e.Deck == 'A':
            img = plt.imread("RubyFPSODeckA.jpg")
        elif (e.Deck == 'B') or (e.Deck == 'C'):
            img = plt.imread("RubyFPSODeckB.jpg")
        elif e.Deck == 'Hull':
            # print('Hull deck?')
            img = plt.imread("RubyFPSOHullDeck.jpg")
        
        if (e.EarlyPoolFire != None) or (e.LatePoolFire != None):
            fig,ax = plt.subplots(3,1,figsize=(8,7.5))
            for axid in [0,1,2]:
                ax[axid].imshow(img, extent=[-11, 252, -32.43, 50.19])
        else:
            fig,ax = plt.subplots(2,1,figsize=(8,5))
            for axid in [0,1]:
                ax[axid].imshow(img, extent=[-11, 252, -32.43, 50.19])

        #Plot Dispersion - LFL
        if (e.Dispersion != None) and (e.Dispersion.DMin != ' '):
            DMin = e.Dispersion.DMin
            DMax = e.Dispersion.DMax
            W = e.Dispersion.W    
            AxisLong = (DMax - DMin)
            AxisShort = W
            DX = (DMax + DMin)*0.5            
            cloud = Ellipse((X-XDir*DX,Y),AxisLong,AxisShort,0)
            # ax.add_artist(cloud)
            ax[0].add_patch(cloud)
            cloud.set_alpha(0.9)
            cloud.set_facecolor('Magenta')
            cloud.set(label='LFL')
        
        #Plot Dispersion - Flash Fire
        if (e.Dispersion != None) and (e.Dispersion.FFMaxDistance != ' '):        
            FFMax = e.Dispersion.FFMaxDistance
            FFW = e.Dispersion.FFWidth
            AxisLong = FFMax
            AxisShort = FFW
            
            cloud = Ellipse((X-XDir*AxisLong*0.5,Y),AxisLong,AxisShort,0)
            # ax.add_artist(cloud)
            ax[0].add_patch(cloud)
            cloud.set_alpha(0.3)
            cloud.set_facecolor('Yellow')
            cloud.set(label='Flash (50\% LFL)')
            ax[0].legend()
        ax[0].set_title(Key)
        


        #Plot Jet Fire
        if e.JetFire != None:
            jfl = e.JetFire.Length
            jfw = 0.12*jfl

            #5, 12, 37.5 kW/m2 contour     
            alpha = 0.8*0.8*0.8     
            w_rad = {5:1.75, 12.5:1.45, 37.5:1.2}
            for rad in [5, 12.5, 37.5]:
                # jfl_rad = np.sqrt(e.JetFire.SEP/rad)+jfl
                jfl_rad = jfl * w_rad[rad]
                jfw_rad = 0.12*jfl_rad

                #Radiation as an ellipse ... the max width is not at the flame tip but at the mid of the flame length
                cloud = Ellipse((X-XDir*jfl_rad*0.5,Y),jfl_rad,jfw_rad,0)                
                ax[1].add_patch(cloud)
                cloud.set_alpha(alpha)
                alpha /= 0.8
                cloud.set_facecolor('Green')
                cloud.set(label="Jet-{:4.1f}kW/m2".format(rad))                
            
            #jet fire shape
            JF = np.array([[X,Y],[X-XDir*jfl,Y+0.5*jfw],[X-XDir*jfl,Y-0.5*jfw]])
            jfpatch = plt.Polygon(JF,color='red')
            ax[1].add_patch(jfpatch)
            jfpatch.set(label="Jet({:3.0f}kW/m2)".format(e.JetFire.SEP))
            ax[1].legend()

        #Plot Early Pool Fire        
        if e.EarlyPoolFire != None:
            EPFD = e.EarlyPoolFire.Diameter
            if e.LatePoolFire != None:            
                LPFD = e.LatePoolFire.Diameter
                PFD = max(EPFD,LPFD)
                SEP = max(e.EarlyPoolFire.SEP,e.LatePoolFire.SEP)
            else:
                PFD = EPFD
                SEP = e.EarlyPoolFire.SEP
            #5, 12, 37.5 kW/m2 contour     
            alpha = 0.8*0.8*0.8     
            for rad in [5, 12.5, 37.5]:
                
                #Total length i.e. diameter is to be calculated
                # D_rad = 2*(np.sqrt(SEP/rad)+PFD/2)
                if PFD > 100:
                    if rad == 37.5:
                        D_rad = 1.*PFD 
                    elif rad == 12.5:
                        D_rad = 1.*PFD
                    elif rad == 5:
                        D_rad = 77/50*PFD                    
                else:
                    if rad == 37.5:
                        D_rad = 1.*PFD 
                    elif rad == 12.5:
                        D_rad = 2*f12(PFD)
                    elif rad == 5:
                        D_rad = 2*f5(PFD)
                # print (rad,D_rad)

                #Radiation as an ellipse ... the max width is not at the flame tip but at the mid of the flame length                
                cloud = Ellipse((X,Y),D_rad,D_rad,0)
                ax[2].add_patch(cloud)
                cloud.set_alpha(alpha)
                alpha /= 0.8
                cloud.set_facecolor('blue')
                cloud.set(label="Pool-{:4.1f}kW/m2".format(rad))                

        if e.LatePoolFire != None:            
            LPF = Ellipse((X,Y),LPFD,LPFD,0)            
            ax[2].add_patch(LPF)
            LPF.set_alpha(1)
            LPF.set_facecolor('cyan')
            LPF.set(label="'Late Pool-{:4.1f}kW/m2".format(e.LatePoolFire.SEP))
        if e.EarlyPoolFire != None:            
            EPF = Ellipse((X,Y),EPFD,EPFD,0)            
            ax[2].add_patch(EPF)
            EPF.set_alpha(0.5)
            EPF.set_facecolor('blue')
            EPF.set(label="'Early Pool-{:4.1f}kW/m2".format(e.EarlyPoolFire.SEP))
        if (e.EarlyPoolFire != None) or (e.LatePoolFire != None):
            ax[2].legend()
        
        
        # ax.clabel(cs2,fmt='%.1e',colors='k',fontsize=14)
        # ax.set_aspect('equal')
        # ax.xaxis.set_major_locator(plt.FixedLocator([120, 141, 168, 193]))
        # ax.xaxis.set_major_formatter(plt.FixedFormatter(['2/3','3/4','4/5','S05']))
        # ax.yaxis.set_major_locator(plt.FixedLocator([-27, -3.1, 3.1, 27]))
        # ax.yaxis.set_major_formatter(plt.FixedFormatter(['ER_S','Tray_S','Tray_P','ER_P']))

        
        # fn = 'C\\ProcessArea\\C_'+slugify(Key)+'.png'
        fn = 'C\\'+folder+'\\C_'+slugify(Key)+'_2.png'

        fig.savefig(fn) 
        # plt.show()
        plt.close()

def Figure(paragraph):
    run = run = paragraph.add_run()
    r = run._r
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    r.append(fldChar)
    instrText = OxmlElement('w:instrText')
    instrText.text = ' STYLEREF 1 \s '
    r.append(instrText)
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    r.append(fldChar)
    instrText = OxmlElement('w:instrText')
    instrText.text = '.'
    r.append(instrText)    
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    r.append(fldChar)
    instrText = OxmlElement('w:instrText')
    instrText.text = ' SEQ Figure \* ARABIC \s 1'
    r.append(instrText)
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    r.append(fldChar)

def Table(paragraph):
    run = run = paragraph.add_run()
    r = run._r
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    r.append(fldChar)
    instrText = OxmlElement('w:instrText')
    instrText.text = ' STYLEREF 1 \s '
    r.append(instrText)
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    r.append(fldChar)
    instrText = OxmlElement('w:instrText')
    instrText.text = '.'
    r.append(instrText)       
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    r.append(fldChar)
    instrText = OxmlElement('w:instrText')
    instrText.text = ' SEQ Table \* ARABIC'
    r.append(instrText)
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    r.append(fldChar)

from docx import Document
from docx.shared import Cm, RGBColor, Pt
from docx.text.run import Font, Run
from docx.shared import Inches, Mm
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
# from docx.dml.color import ColorFormat

document = Document('form.docx')

fontn = document.styles['Normal'].font
fontn.size = Pt(9)
fontn.name = 'Frutiger LT 45 Light'
fontn.color.rgb = RGBColor(31, 73, 125)

# document.add_heading('PHAST Analysis - ' + folder, level=0)

ISs = []
for e in lEvent:
    pvis,hole,weather = e.Key.split("\\")
    if not (pvis in ISs):
        ISs.append(pvis)

for IS in ISs:
    document.add_heading(IS, level=1)
    for e in lEvent:
        pvis,hole,weather = e.Key.split("\\")
        if IS == pvis:                        
            fn = 'C\\'+folder+'\\C_'+slugify(e.Key)+'_2.png'
            document.add_picture(fn, width=Cm(15))            

            p = document.add_paragraph('Figure',style='FigureCaption')
            Figure(p)
            p.add_run(" " + pvis + " Hole: " + hole + " Weather: " + weather)            
    document.add_page_break()
document.save('C_'+folder+'_Rev.2.docx')
